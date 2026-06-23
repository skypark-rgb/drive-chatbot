from io import BytesIO

from docx import Document as WordDocument

from models.document import Document


class WordParser:
    def __init__(self, drive_client):
        self.drive_client = drive_client

    def parse(self, file_metadata):
        docx_bytes = self.drive_client.download_file(
            file_metadata.id
        )

        document = WordDocument(BytesIO(docx_bytes))

        text = []

        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)

        return Document(
            file_id=file_metadata.id,
            name=file_metadata.name,
            mime_type=file_metadata.mime_type,
            text="\n".join(text),
        )